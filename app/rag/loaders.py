"""Document ingestion — file bytes to normalised plain text.

Dispatch is on extension first, MIME second. Every loader returns
(text, page_map) where page_map lets citations point back to a page number
for paginated formats.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from app.core.errors import UnsupportedFileType

SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".markdown", ".html", ".htm", ".csv"}

_WS = re.compile(r"[ \t\x0b\f\r]+")
_BLANKS = re.compile(r"\n{3,}")


def normalise(text: str) -> str:
    """Collapse whitespace and de-hyphenate line-wrapped words."""
    text = text.replace("\u00ad", "")  # soft hyphen
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)  # join hyphenated line breaks
    text = _WS.sub(" ", text)
    text = _BLANKS.sub("\n\n", text)
    return text.strip()


def _load_pdf(data: bytes) -> tuple[str, list[int]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    page_map: list[int] = []
    for page_no, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        parts.append(text)
        page_map.append(page_no)
    return normalise("\n\n".join(parts)), page_map


def _load_docx(data: bytes) -> tuple[str, list[int]]:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    blocks = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return normalise("\n\n".join(blocks)), []


def _load_html(data: bytes) -> tuple[str, list[int]]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data.decode("utf-8", errors="replace"), "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return normalise(soup.get_text("\n")), []


def _load_text(data: bytes) -> tuple[str, list[int]]:
    return normalise(data.decode("utf-8", errors="replace")), []


_DISPATCH = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".html": _load_html,
    ".htm": _load_html,
    ".txt": _load_text,
    ".md": _load_text,
    ".markdown": _load_text,
    ".csv": _load_text,
}


def load(filename: str, data: bytes) -> tuple[str, list[int]]:
    ext = Path(filename).suffix.lower()
    loader = _DISPATCH.get(ext)
    if loader is None:
        raise UnsupportedFileType(
            f"'{ext or filename}' is not supported. Accepted: {', '.join(sorted(SUPPORTED))}"
        )
    return loader(data)
